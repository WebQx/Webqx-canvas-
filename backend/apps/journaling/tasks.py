from celery import shared_task
from django.utils import timezone
from datetime import timedelta
import os
import json

from .models import JournalEntry, JournalExport
from .nlp_service import JournalNLPService


@shared_task
def process_journal_entry_nlp(entry_id):
    """Process NLP analysis for a journal entry"""
    try:
        entry = JournalEntry.objects.get(id=entry_id)
        
        nlp_service = JournalNLPService()
        analysis = nlp_service.analyze_entry(entry)
        
        # Update entry with analysis results
        entry.sentiment_score = analysis['sentiment_score']
        entry.sentiment_label = analysis['sentiment_label']
        entry.keywords = analysis['keywords']
        entry.entities = analysis['entities']
        entry.topics = analysis['topics']
        entry.urgency_score = analysis['urgency_score']
        entry.clinical_flags = analysis['clinical_flags']
        entry.save()
        
        return f"NLP analysis completed for entry {entry_id}"
        
    except JournalEntry.DoesNotExist:
        return f"Journal entry {entry_id} not found"
    except Exception as e:
        return f"Error processing NLP for entry {entry_id}: {str(e)}"


@shared_task
def generate_journal_export(export_id):
    """Generate journal export file"""
    try:
        export = JournalExport.objects.get(id=export_id)
        
        # Get journal entries for the date range
        entries = JournalEntry.objects.filter(
            user=export.user,
            created_at__date__gte=export.date_range_start,
            created_at__date__lte=export.date_range_end
        )
        
        # Apply filters
        if not export.include_private:
            entries = entries.filter(is_private=False)
        
        if export.entry_types:
            entries = entries.filter(entry_type__in=export.entry_types)
        
        # Generate file based on format
        if export.export_format == 'json':
            file_path = _generate_json_export(export, entries)
        elif export.export_format == 'csv':
            file_path = _generate_csv_export(export, entries)
        elif export.export_format == 'pdf':
            file_path = _generate_pdf_export(export, entries)
        elif export.export_format == 'docx':
            file_path = _generate_docx_export(export, entries)
        else:
            raise ValueError(f"Unsupported export format: {export.export_format}")
        
        # Update export record
        export.file_path = file_path
        export.file_size_bytes = os.path.getsize(file_path)
        export.is_complete = True
        export.expires_at = timezone.now() + timedelta(days=7)  # Expire in 7 days
        export.save()
        
        return f"Export {export_id} completed successfully"
        
    except JournalExport.DoesNotExist:
        return f"Export {export_id} not found"
    except Exception as e:
        # Update export with error
        try:
            export = JournalExport.objects.get(id=export_id)
            export.error_message = str(e)
            export.save()
        except:
            pass
        return f"Error generating export {export_id}: {str(e)}"


def _generate_json_export(export, entries):
    """Generate JSON export"""
    from django.core import serializers
    import json
    
    data = {
        'export_info': {
            'user': export.user.username,
            'created_at': export.created_at.isoformat(),
            'date_range': {
                'start': export.date_range_start.isoformat(),
                'end': export.date_range_end.isoformat()
            },
            'total_entries': entries.count()
        },
        'entries': []
    }
    
    for entry in entries:
        entry_data = {
            'id': entry.id,
            'title': entry.title,
            'content': entry.content,
            'entry_type': entry.entry_type,
            'mood_rating': entry.mood_rating,
            'pain_level': entry.pain_level,
            'created_at': entry.created_at.isoformat(),
            'sentiment_score': entry.sentiment_score,
            'sentiment_label': entry.sentiment_label,
            'keywords': entry.keywords,
            'topics': entry.topics,
            'word_count': entry.word_count
        }
        data['entries'].append(entry_data)
    
    # Save to file
    filename = f"journal_export_{export.id}_{timezone.now().timestamp()}.json"
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    
    with open(file_path, 'w') as f:
        json.dump(data, f, indent=2)
    
    return file_path


def _generate_csv_export(export, entries):
    """Generate CSV export"""
    import csv
    
    filename = f"journal_export_{export.id}_{timezone.now().timestamp()}.csv"
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    
    with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = [
            'id', 'title', 'content', 'entry_type', 'mood_rating',
            'pain_level', 'created_at', 'sentiment_score', 'sentiment_label',
            'keywords', 'topics', 'word_count'
        ]
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        
        writer.writeheader()
        for entry in entries:
            writer.writerow({
                'id': entry.id,
                'title': entry.title,
                'content': entry.content,
                'entry_type': entry.entry_type,
                'mood_rating': entry.mood_rating,
                'pain_level': entry.pain_level,
                'created_at': entry.created_at.isoformat(),
                'sentiment_score': entry.sentiment_score,
                'sentiment_label': entry.sentiment_label,
                'keywords': ', '.join(entry.keywords) if entry.keywords else '',
                'topics': ', '.join(entry.topics) if entry.topics else '',
                'word_count': entry.word_count
            })
    
    return file_path


def _generate_pdf_export(export, entries):
    """Generate PDF export"""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    
    filename = f"journal_export_{export.id}_{timezone.now().timestamp()}.pdf"
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title = Paragraph(f"Journal Export for {export.user.full_name}", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    # Export info
    info = Paragraph(
        f"Export Date: {export.created_at.strftime('%B %d, %Y')}<br/>"
        f"Date Range: {export.date_range_start} to {export.date_range_end}<br/>"
        f"Total Entries: {entries.count()}",
        styles['Normal']
    )
    story.append(info)
    story.append(Spacer(1, 24))
    
    # Entries
    for entry in entries:
        entry_title = Paragraph(
            f"{entry.title or 'Journal Entry'} - {entry.created_at.strftime('%B %d, %Y')}",
            styles['Heading2']
        )
        story.append(entry_title)
        
        content = Paragraph(entry.content, styles['Normal'])
        story.append(content)
        story.append(Spacer(1, 12))
        
        if entry.mood_rating:
            mood = Paragraph(f"Mood: {entry.mood_rating}/5", styles['Normal'])
            story.append(mood)
        
        if entry.pain_level:
            pain = Paragraph(f"Pain Level: {entry.pain_level}/10", styles['Normal'])
            story.append(pain)
        
        story.append(Spacer(1, 24))
    
    doc.build(story)
    return file_path


def _generate_docx_export(export, entries):
    """Generate DOCX export"""
    from docx import Document
    
    filename = f"journal_export_{export.id}_{timezone.now().timestamp()}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Journal Export for {export.user.full_name}', 0)
    
    # Export info
    doc.add_paragraph(f'Export Date: {export.created_at.strftime("%B %d, %Y")}')
    doc.add_paragraph(f'Date Range: {export.date_range_start} to {export.date_range_end}')
    doc.add_paragraph(f'Total Entries: {entries.count()}')
    doc.add_paragraph('')
    
    # Entries
    for entry in entries:
        entry_heading = doc.add_heading(
            f'{entry.title or "Journal Entry"} - {entry.created_at.strftime("%B %d, %Y")}',
            level=1
        )
        
        doc.add_paragraph(entry.content)
        
        if entry.mood_rating:
            doc.add_paragraph(f'Mood: {entry.mood_rating}/5')
        
        if entry.pain_level:
            doc.add_paragraph(f'Pain Level: {entry.pain_level}/10')
        
        doc.add_paragraph('')  # Add space between entries
    
    doc.save(file_path)
    return file_path


@shared_task
def cleanup_expired_exports():
    """Clean up expired journal exports"""
    expired_exports = JournalExport.objects.filter(
        expires_at__lt=timezone.now(),
        is_complete=True
    )
    
    deleted_count = 0
    for export in expired_exports:
        try:
            if export.file_path and os.path.exists(export.file_path):
                os.remove(export.file_path)
            export.delete()
            deleted_count += 1
        except Exception as e:
            print(f"Error deleting export {export.id}: {str(e)}")
    
    return f"Cleaned up {deleted_count} expired exports"